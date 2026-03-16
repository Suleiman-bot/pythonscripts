from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, StaleElementReferenceException
import time
import re
import urllib3
from config import ACCESS_TOKEN, TARGET_DATE

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    Document = None

# === Construct magic login link (same approach as export.py) ===
magic_link = f"https://www.cv-prod-euwest-2.arista.io/api/v1/oauth?invitation={ACCESS_TOKEN}"

def generate_active_for_target_date_local(target_date_str=None):
    """
    Returns 'active' timestamp in milliseconds for the target date's 00:00:00
    in local timezone (Africa/Lagos, UTC+1)
    
    Args:
        target_date_str: Date string in format "M/D/YYYY". If None, uses TARGET_DATE from config
    """
    from datetime import datetime
    import pytz
    if target_date_str is None:
        target_date_str = TARGET_DATE
        
    local_tz = pytz.timezone("Africa/Lagos")
    
    # Parse the target date string in format "M/D/YYYY"
    target_date = datetime.strptime(target_date_str, "%m/%d/%Y")
    
    # Localize to timezone without adding extra hours
    target_date_local = local_tz.localize(target_date)
    
    # Convert to UTC timestamp in milliseconds
    active_ts = int(target_date_local.timestamp() * 1000)
    return active_ts

ACTIVE = generate_active_for_target_date_local()
FROM_OFFSET = 1000
TO_OFFSET = 86400000

# --- Data Usage Links ---
DATA_USAGE_LINKS = {
    "Table1": {
        "Glo": {
            "Inflow": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/traffic-flows/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&trafficFlowFilters=%7B%22bidirectional%22%3Afalse%2C%22device_ids%22%3A%5B%5D%2C%22latencies%22%3A%5B%5D%2C%22exclude%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%2C%22include%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%22Ethernet1%22%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%7D",
            "Outflow": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/traffic-flows/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&trafficFlowFilters=%7B%22bidirectional%22%3Afalse%2C%22device_ids%22%3A%5B%5D%2C%22latencies%22%3A%5B%5D%2C%22exclude%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%2C%22include%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%22Ethernet1%22%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%7D",
            "Total": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/traffic-flows/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&trafficFlowFilters=%7B%22bidirectional%22%3Afalse%2C%22device_ids%22%3A%5B%5D%2C%22latencies%22%3A%5B%5D%2C%22exclude%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%2C%22include%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%22Ethernet1%22%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%7D"
        },
        "Dolphin Sec.": {
            "Inflow": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/traffic-flows/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&trafficFlowFilters=%7B%22bidirectional%22%3Afalse%2C%22device_ids%22%3A%5B%5D%2C%22latencies%22%3A%5B%5D%2C%22exclude%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%2C%22include%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%22Ethernet2%22%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%7D",
            "Outflow": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/traffic-flows/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&trafficFlowFilters=%7B%22bidirectional%22%3Afalse%2C%22device_ids%22%3A%5B%5D%2C%22latencies%22%3A%5B%5D%2C%22exclude%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%2C%22include%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%22Ethernet2%22%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%7D",
            "Total": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/traffic-flows/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&trafficFlowFilters=%7B%22bidirectional%22%3Afalse%2C%22device_ids%22%3A%5B%5D%2C%22latencies%22%3A%5B%5D%2C%22exclude%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%2C%22include%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%22Ethernet2%22%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%7D"
        },
        "Dolphin Pri.": {
            "Inflow": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/traffic-flows/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&trafficFlowFilters=%7B%22bidirectional%22%3Afalse%2C%22device_ids%22%3A%5B%5D%2C%22latencies%22%3A%5B%5D%2C%22exclude%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%2C%22include%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%22Ethernet3%22%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%7D",
            "Outflow": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/traffic-flows/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&trafficFlowFilters=%7B%22bidirectional%22%3Afalse%2C%22device_ids%22%3A%5B%5D%2C%22latencies%22%3A%5B%5D%2C%22exclude%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%2C%22include%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%22Ethernet3%22%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%7D",
            "Total": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/traffic-flows/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&trafficFlowFilters=%7B%22bidirectional%22%3Afalse%2C%22device_ids%22%3A%5B%5D%2C%22latencies%22%3A%5B%5D%2C%22exclude%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%2C%22include%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%22Ethernet3%22%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%7D"
        },
        "Total Data Usage:": {
            "Inflow": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/traffic-flows/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&trafficFlowFilters=%7B%22bidirectional%22%3Afalse%2C%22device_ids%22%3A%5B%5D%2C%22latencies%22%3A%5B%5D%2C%22exclude%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%2C%22include%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%22Ethernet1%22%2C%22Ethernet2%22%2C%22Ethernet3%22%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%7D",
            "Outflow": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/traffic-flows/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&trafficFlowFilters=%7B%22bidirectional%22%3Afalse%2C%22device_ids%22%3A%5B%5D%2C%22latencies%22%3A%5B%5D%2C%22exclude%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%2C%22include%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%22Ethernet1%22%2C%22Ethernet2%22%2C%22Ethernet3%22%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%7D",
            "Total": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/traffic-flows/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&trafficFlowFilters=%7B%22bidirectional%22%3Afalse%2C%22device_ids%22%3A%5B%5D%2C%22latencies%22%3A%5B%5D%2C%22exclude%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%2C%22include%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%22Ethernet1%22%2C%22Ethernet2%22%2C%22Ethernet3%22%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%7D"
        }
    },
    "Table2": {
        "HQ Campus": {
            "Inflow": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/traffic-flows/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&trafficFlowFilters=%7B%22bidirectional%22%3Afalse%2C%22device_ids%22%3A%5B%5D%2C%22latencies%22%3A%5B%5D%2C%22exclude%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%2C%22include%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%22Ethernet4%22%2C%22Ethernet5%22%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%7D",
            "Outflow": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/traffic-flows/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&trafficFlowFilters=%7B%22bidirectional%22%3Afalse%2C%22device_ids%22%3A%5B%5D%2C%22latencies%22%3A%5B%5D%2C%22exclude%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%2C%22include%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%22Ethernet4%22%2C%22Ethernet5%22%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%7D",
            "Total": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/traffic-flows/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&trafficFlowFilters=%7B%22bidirectional%22%3Afalse%2C%22device_ids%22%3A%5B%5D%2C%22latencies%22%3A%5B%5D%2C%22exclude%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%2C%22include%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%22Ethernet4%22%2C%22Ethernet5%22%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%7D"
        },
        "Foundry Customers": {
            "Inflow": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/traffic-flows/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&trafficFlowFilters=%7B%22bidirectional%22%3Afalse%2C%22device_ids%22%3A%5B%5D%2C%22latencies%22%3A%5B%5D%2C%22exclude%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%2C%22include%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%22Ethernet49%2F1%22%2C%22Ethernet50%2F1%22%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%7D",
            "Outflow": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/traffic-flows/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&trafficFlowFilters=%7B%22bidirectional%22%3Afalse%2C%22device_ids%22%3A%5B%5D%2C%22latencies%22%3A%5B%5D%2C%22exclude%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%2C%22include%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%22Ethernet49%2F1%22%2C%22Ethernet50%2F1%22%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%7D",
            "Total": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/traffic-flows/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&trafficFlowFilters=%7B%22bidirectional%22%3Afalse%2C%22device_ids%22%3A%5B%5D%2C%22latencies%22%3A%5B%5D%2C%22exclude%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%2C%22include%22%3A%7B%22avt_names%22%3A%5B%5D%2C%22dst_applications%22%3A%5B%5D%2C%22dst_ips%22%3A%5B%5D%2C%22dst_ports%22%3A%5B%5D%2C%22dpi_applications%22%3A%5B%5D%2C%22dps_path_groups%22%3A%5B%5D%2C%22dps_path_ids%22%3A%5B%5D%2C%22dps_path_local_interfaces%22%3A%5B%5D%2C%22dps_path_remote_device_ids%22%3A%5B%5D%2C%22dps_path_remote_interfaces%22%3A%5B%5D%2C%22egress_interfaces%22%3A%5B%5D%2C%22ingress_interfaces%22%3A%5B%5D%2C%22interfaces%22%3A%5B%22Ethernet49%2F1%22%2C%22Ethernet50%2F1%22%5D%2C%22ip_class_of_service%22%3A%5B%5D%2C%22protocols%22%3A%5B%5D%2C%22src_applications%22%3A%5B%5D%2C%22src_ips%22%3A%5B%5D%2C%22src_ports%22%3A%5B%5D%2C%22tunnel_ids%22%3A%5B%5D%2C%22user_identity%22%3A%5B%5D%2C%22vrf_name%22%3A%5B%5D%2C%22vlan_id%22%3A%5B%5D%2C%22bgp_dst_as%22%3A%5B%5D%2C%22bgp_dst_peer_as%22%3A%5B%5D%2C%22bgp_src_as%22%3A%5B%5D%2C%22bgp_src_peer_as%22%3A%5B%5D%7D%7D"
        }
    }
}

DOCX_PATH = r"C:\Users\SuleimanAbdulsalam\Downloads\Data-Usage.docx"


# Utility: extract data usage value
def extract_data_usage(driver, timeout=12):
    start = time.time()
    while time.time() - start < timeout:
        try:
            # Find element containing "Data Usage"
            data_usage_el = driver.find_element(By.XPATH, "//*[contains(translate(text(),'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'data usage')]")
            # Get the text and extract the value next to it
            text = data_usage_el.text
            # Look for pattern like "Data Usage: 1.5 GB" or similar
            match = re.search(r'Data Usage[:\s]*([0-9]+\.?[0-9]*\s*[GM]B)', text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
            # If not in the element, check parent or sibling
            parent = data_usage_el.find_element(By.XPATH, "..")
            parent_text = parent.text
            match = re.search(r'Data Usage[:\s]*([0-9]+\.?[0-9]*\s*[GM]B)', parent_text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        except Exception:
            pass
        time.sleep(0.5)
    return "-"


# Collect data usage
def collect_data_usage(driver, links):
    results = {}
    for table, table_data in links.items():
        results[table] = {}
        for name, name_data in table_data.items():
            results[table][name] = {}
            for direction, url in name_data.items():
                try:
                    driver.get(url)
                    time.sleep(0.5)
                    print(f"Opening {table} {name} {direction} URL...")
                    value = extract_data_usage(driver)
                    results[table][name][direction] = value
                    print(f"Result {table} {name} {direction} -> '{value}'")
                except Exception as e:
                    print(f"Error collecting {table} {name} {direction}: {e}")
                    results[table][name][direction] = "-"
    return results


# Save DOCX
def save_docx(results, path):
    if Document is None:
        print("python-docx not installed. Install with: pip install python-docx")
        return
    doc = Document()
    for table_name, table_data in results.items():
        if table_name == "Table1":
            doc.add_heading("NETWORK FLOW DATA/DATA USAGE:", level=2)
            col0 = 'ISP Link'
            col1 = 'Inflow (Ingress Traffic)'
            col2 = 'Outflow (Egress Traffic)'
            col3 = 'Total'
        elif table_name == "Table2":
            doc.add_heading("Consumption Breakdown", level=2)
            col0 = 'Consumption Breakdown'
            col1 = 'Inflow (Ingress Traffic)'
            col2 = 'Outflow (Egress Traffic)'
            col3 = 'Total Data Usage'
        else:
            doc.add_heading(table_name, level=2)
            col0 = 'ISP'
            col1 = 'Inflow'
            col2 = 'Outflow'
            col3 = 'Total'

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        # Header
        hdr = table.rows[0].cells
        hdr[0].text = col0
        hdr[1].text = col1
        hdr[2].text = col2
        hdr[3].text = col3
        # Bold headers
        for cell in hdr:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        # Rows
        for name in table_data:
            row_cells = table.add_row().cells
            row_cells[0].text = name
            row_cells[1].text = table_data[name].get('Inflow', '-')
            row_cells[2].text = table_data[name].get('Outflow', '-')
            row_cells[3].text = table_data[name].get('Total', '-')
        doc.add_paragraph("")  # Space between tables
    doc.save(path)
    print(f"Saved DOCX to: {path}")


def run():
    print("Opening Selenium with Edge profile...")
    options = webdriver.EdgeOptions()
    options.add_argument("--start-maximized")
    options.add_argument(r"--user-data-dir=C:\Users\SuleimanAbdulsalam\AppData\Local\Microsoft\Edge\User Data\SeleniumProfile")
    options.add_argument("profile-directory=Default")
    driver = webdriver.Edge(options=options)

    try:
        print("Opening CVaaS login link...")
        driver.get(magic_link)
        try:
            WebDriverWait(driver, 30).until(lambda d: d.current_url != magic_link or d.execute_script("return document.readyState") == "complete")
        except TimeoutException:
            pass

        print("Collecting data usage...")
        results = collect_data_usage(driver, DATA_USAGE_LINKS)

        # Save DOCX
        save_docx(results, DOCX_PATH)
        print(f"Saved to: {DOCX_PATH}")
    finally:
        driver.quit()


if __name__ == "__main__":
    run()