from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, StaleElementReferenceException
import time
import re
import csv
from datetime import datetime, timedelta
import pytz
import urllib3
from config import ACCESS_TOKEN, TARGET_DATE
import json

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    Document = None

# === Construct magic login link (same approach as statistics.py) ===
magic_link = f"https://www.cv-prod-euwest-2.arista.io/api/v1/oauth?invitation={ACCESS_TOKEN}"


def generate_active_for_target_date_local(target_date_str=None):
    """
    Returns 'active' timestamp in milliseconds for the target date's 00:00:00
    in local timezone (Africa/Lagos, UTC+1)
    
    Args:
        target_date_str: Date string in format "M/D/YYYY". If None, uses TARGET_DATE from config
    """
    if target_date_str is None:
        target_date_str = TARGET_DATE
        
    local_tz = pytz.timezone("Africa/Lagos")
    
    # Parse the target date string in format "M/D/YYYY"
    target_date = datetime.strptime(target_date_str, "%m/%d/%Y")
    
    # Localize to timezone without adding extra hours
    target_date_local = local_tz.localize(target_date)
    
    # Convert to UTC timestamp in milliseconds
    active_ts = int(target_date_local.timestamp() * 1000)
    return active_ts

ACTIVE = generate_active_for_target_date_local()
FROM_OFFSET = 1000
TO_OFFSET = 86400000
# === End of active generation function ===


# Device Health Links with Device IDs and proper labels
DEVICE_HEALTH_LINKS = {
    "KASI-LOS5-R201-BG01": {
        "device_id": "JPE20050335",
        "memory": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/processes/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22collisionWarning%22%3Afalse%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22DEVICE_MEMORY_USAGE_PERCENTAGE%22%2C%22metricParams%22%3A%7B%22memberId%22%3A%221%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22deviceId%22%3A%22JPE20050335%22%7D%7D&modalPanel=STATISTICS",
        "cpu": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/processes/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22collisionWarning%22%3Afalse%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22DEVICE_CPU%22%2C%22metricParams%22%3A%7B%22memberId%22%3A%221%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22deviceId%22%3A%22JPE20050335%22%7D%7D&modalPanel=STATISTICS",
    },
    "KASI-LOS5-R201-EOR1": {
        "device_id": "SSJ17243133",
        "memory": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/processes/SSJ17243133?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22collisionWarning%22%3Afalse%2C%22datasetId%22%3A%22SSJ17243133%22%2C%22metricKey%22%3A%22DEVICE_MEMORY_USAGE_PERCENTAGE%22%2C%22metricParams%22%3A%7B%22memberId%22%3A%221%22%2C%22datasetId%22%3A%22SSJ17243133%22%2C%22deviceId%22%3A%22SSJ17243133%22%7D%7D&modalPanel=STATISTICS",
        "cpu": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/processes/SSJ17243133?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22collisionWarning%22%3Afalse%2C%22datasetId%22%3A%22SSJ17243133%22%2C%22metricKey%22%3A%22DEVICE_CPU%22%2C%22metricParams%22%3A%7B%22memberId%22%3A%221%22%2C%22datasetId%22%3A%22SSJ17243133%22%2C%22deviceId%22%3A%22SSJ17243133%22%7D%7D&modalPanel=STATISTICS",
    },
    "KASI-LOS5-R101-TOR2": {
        "device_id": "SSJ17243158",
        "memory": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/processes/SSJ17243158?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22collisionWarning%22%3Afalse%2C%22datasetId%22%3A%22SSJ17243158%22%2C%22metricKey%22%3A%22DEVICE_MEMORY_USAGE_PERCENTAGE%22%2C%22metricParams%22%3A%7B%22memberId%22%3A%221%22%2C%22datasetId%22%3A%22SSJ17243158%22%2C%22deviceId%22%3A%22SSJ17243158%22%7D%7D&modalPanel=STATISTICS",
        "cpu": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/processes/SSJ17243158?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22collisionWarning%22%3Afalse%2C%22datasetId%22%3A%22SSJ17243158%22%2C%22metricKey%22%3A%22DEVICE_CPU%22%2C%22metricParams%22%3A%7B%22memberId%22%3A%221%22%2C%22datasetId%22%3A%22SSJ17243158%22%2C%22deviceId%22%3A%22SSJ17243158%22%7D%7D&modalPanel=STATISTICS",
    },
    "KASI-LOS5-R201-EOR2": {
        "device_id": "JPE17015583",
        "memory": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/processes/JPE17015583?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22collisionWarning%22%3Afalse%2C%22datasetId%22%3A%22JPE17015583%22%2C%22metricKey%22%3A%22DEVICE_MEMORY_USAGE_PERCENTAGE%22%2C%22metricParams%22%3A%7B%22memberId%22%3A%221%22%2C%22datasetId%22%3A%22JPE17015583%22%2C%22deviceId%22%3A%22JPE17015583%22%7D%7D&modalPanel=STATISTICS",
        "cpu": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/processes/JPE17015583?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22collisionWarning%22%3Afalse%2C%22datasetId%22%3A%22JPE17015583%22%2C%22metricKey%22%3A%22DEVICE_CPU%22%2C%22metricParams%22%3A%7B%22memberId%22%3A%221%22%2C%22datasetId%22%3A%22JPE17015583%22%2C%22deviceId%22%3A%22JPE17015583%22%7D%7D&modalPanel=STATISTICS",
    },
    "KASI-LOS5-R201-TOR1": {
        "device_id": "HSH16170211",
        "memory": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/processes/HSH16170211?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22collisionWarning%22%3Afalse%2C%22datasetId%22%3A%22HSH16170211%22%2C%22metricKey%22%3A%22DEVICE_MEMORY_USAGE_PERCENTAGE%22%2C%22metricParams%22%3A%7B%22memberId%22%3A%221%22%2C%22datasetId%22%3A%22HSH16170211%22%2C%22deviceId%22%3A%22HSH16170211%22%7D%7D&modalPanel=STATISTICS",
        "cpu": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/processes/HSH16170211?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22collisionWarning%22%3Afalse%2C%22datasetId%22%3A%22HSH16170211%22%2C%22metricKey%22%3A%22DEVICE_CPU%22%2C%22metricParams%22%3A%7B%22memberId%22%3A%221%22%2C%22datasetId%22%3A%22HSH16170211%22%2C%22deviceId%22%3A%22HSH16170211%22%7D%7D&modalPanel=STATISTICS",
    },
    "KASI-HQ-FLC1-TOR1": {
        "device_id": "WTW22430298",
        "memory": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/processes/WTW22430298?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22collisionWarning%22%3Afalse%2C%22datasetId%22%3A%22WTW22430298%22%2C%22metricKey%22%3A%22DEVICE_MEMORY_USAGE_PERCENTAGE%22%2C%22metricParams%22%3A%7B%22memberId%22%3A%221%22%2C%22datasetId%22%3A%22WTW22430298%22%2C%22deviceId%22%3A%22WTW22430298%22%7D%7D&modalPanel=STATISTICS",
        "cpu": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/processes/WTW22430298?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22collisionWarning%22%3Afalse%2C%22datasetId%22%3A%22WTW22430298%22%2C%22metricKey%22%3A%22DEVICE_CPU%22%2C%22metricParams%22%3A%7B%22memberId%22%3A%221%22%2C%22datasetId%22%3A%22WTW22430298%22%2C%22deviceId%22%3A%22WTW22430298%22%7D%7D&modalPanel=STATISTICS",
    },
    "KASI-HQ-FLD1-TOR2": {
        "device_id": "WTW22421221",
        "memory": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/processes/WTW22421221?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22collisionWarning%22%3Afalse%2C%22datasetId%22%3A%22WTW22421221%22%2C%22metricKey%22%3A%22DEVICE_MEMORY_USAGE_PERCENTAGE%22%2C%22metricParams%22%3A%7B%22memberId%22%3A%221%22%2C%22datasetId%22%3A%22WTW22421221%22%2C%22deviceId%22%3A%22WTW22421221%22%7D%7D&modalPanel=STATISTICS",
        "cpu": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/processes/WTW22421221?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22collisionWarning%22%3Afalse%2C%22datasetId%22%3A%22WTW22421221%22%2C%22metricKey%22%3A%22DEVICE_CPU%22%2C%22metricParams%22%3A%7B%22memberId%22%3A%221%22%2C%22datasetId%22%3A%22WTW22421221%22%2C%22deviceId%22%3A%22WTW22421221%22%7D%7D&modalPanel=STATISTICS",
    },
}

DOCX_DEVICE_HEALTH_PATH = r"C:\Users\SuleimanAbdulsalam\Downloads\device-health-summary.docx"
CSV_DEVICE_HEALTH_PATH = r"C:\Users\SuleimanAbdulsalam\Downloads\device-health-summary.csv"

# Utility: try to parse a number (integers/floats) from a text blob
_num_re = re.compile(r"[-+]?[0-9]+(?:[\,\.][0-9]+)*")

def extract_first_number(text):
    if not text:
        return ""
    m = _num_re.search(text.replace('\u00a0', ' '))
    if not m:
        return ""
    val = m.group(0)
    return val


def find_mean_value_on_page(driver, label="", timeout=12, poll=0.5):
    """
    Find the 'mean' text on the page and extract the numeric value next to it.
    This accesses the statistics page directly on CVaaS.
    """
    start = time.time()

    def _attempt_find():
        # Ensure page loaded
        try:
            WebDriverWait(driver, 3).until(lambda d: d.execute_script("return document.readyState") == "complete")
        except TimeoutException:
            pass

        # Get page text
        try:
            body = driver.find_element(By.TAG_NAME, "body")
            page_txt = body.text
        except Exception:
            return ""

        if not page_txt:
            return ""

        # Look for pattern like "Mean 23.4" or "mean: 23.4" or "Mean 23.4%" etc.
        # Case-insensitive search
        mean_pattern = r"mean\s*[:\-]?\s*([0-9]+(?:[.,][0-9]+)?)"
        match = re.search(mean_pattern, page_txt, re.IGNORECASE)
        
        if match:
            value_str = match.group(1)
            # Normalize decimal separator
            value_str = value_str.replace(',', '.')
            return value_str
        
        return ""

    # Poll until we get value or timeout
    while time.time() - start < timeout:
        mean_val = _attempt_find()
        if mean_val:
            print(f"Found MEAN for {label}: {mean_val}")
            return mean_val
        time.sleep(poll)

    print(f"No mean value found for {label} within {timeout}s")
    return ""


# Open each URL and extract metric values
def collect_device_health_metrics(driver, devices, per_page_timeout=12):
    results = {}
    wait = WebDriverWait(driver, 10)
    
    for device_name, urls in devices.items():
        results[device_name] = {"memory_mean": "", "cpu_mean": ""}
        
        # Memory metric
        try:
            driver.get(urls["memory"])
            time.sleep(1)
            print(f"Opening {device_name} memory URL...")
            memory_mean = find_mean_value_on_page(driver, label=f"{device_name} memory", timeout=per_page_timeout)
            results[device_name]["memory_mean"] = memory_mean
            print(f"Result {device_name} memory -> mean: '{memory_mean}'")
        except Exception as e:
            print(f"Error collecting memory for {device_name}: {e}")
            results[device_name]["memory_mean"] = ""

        # CPU metric
        try:
            driver.get(urls["cpu"])
            time.sleep(1)
            print(f"Opening {device_name} CPU URL...")
            cpu_mean = find_mean_value_on_page(driver, label=f"{device_name} CPU", timeout=per_page_timeout)
            results[device_name]["cpu_mean"] = cpu_mean
            print(f"Result {device_name} CPU -> mean: '{cpu_mean}'")
        except Exception as e:
            print(f"Error collecting CPU for {device_name}: {e}")
            results[device_name]["cpu_mean"] = ""

    return results


# Print summary table
def print_summary_table(title, results):
    print(title)
    print("-" * 70)
    
    for device_name in sorted(results.keys()):
        row = results.get(device_name, {})
        memory = row.get("memory_mean", "") or ""
        cpu = row.get("cpu_mean", "") or ""
        
        memory_str = f"{memory}%" if memory else "N/A"
        cpu_str = f"{cpu}%" if cpu else "N/A"
        
        print(f"{device_name}: Memory Usage: {memory_str}, CPU Utilization: {cpu_str}")
    print()


# Save DOCX
def save_docx_summary(path, title, results):
    if Document is None:
        print("python-docx not installed. Install with: pip install python-docx")
        return
    
    doc = Document()
    doc.add_heading(title, level=2)
    
    # Format: Device Name: Memory Usage: X%, CPU Utilization: Y%
    for device_name in sorted(results.keys()):
        memory = results.get(device_name, {}).get('memory_mean', '') or ''
        cpu = results.get(device_name, {}).get('cpu_mean', '') or ''
        
        memory_str = f"{memory}%" if memory else "N/A"
        cpu_str = f"{cpu}%" if cpu else "N/A"
        
        line = f"{device_name}: Memory Usage: {memory_str}, CPU Utilization: {cpu_str}"
        doc.add_paragraph(line)
    
    doc.save(path)
    print(f"Saved DOCX to: {path}")


# Save CSV
def save_csv_summary(path, title, results):
    with open(path, "w", newline='', encoding='utf-8') as fh:
        w = csv.writer(fh)
        w.writerow([title])
        w.writerow([])  # Blank row for spacing
        for device_name in sorted(results.keys()):
            row = results.get(device_name, {})
            memory = row.get('memory_mean', '') or ''
            cpu = row.get('cpu_mean', '') or ''
            memory_str = f"{memory}%" if memory else ""
            cpu_str = f"{cpu}%" if cpu else ""
            w.writerow([f"{device_name}: Memory Usage: {memory_str}, CPU Utilization: {cpu_str}"])
    print(f"Saved CSV to: {path}")


def run():
    print("🌐 Opening Selenium with Edge profile...")
    options = webdriver.EdgeOptions()
    options.add_argument("--start-maximized")
    options.add_argument(r"--user-data-dir=C:\Users\SuleimanAbdulsalam\AppData\Local\Microsoft\Edge\User Data\SeleniumProfile")
    options.add_argument("profile-directory=Default")
    driver = webdriver.Edge(options=options)

    try:
        print("Opening CVaaS login link...")
        driver.get(magic_link)
        try:
            WebDriverWait(driver, 30).until(lambda d: d.current_url != magic_link or d.execute_script("return document.readyState") == "complete")
        except TimeoutException:
            pass

        print("Collecting device health metrics...")
        health_results = collect_device_health_metrics(driver, DEVICE_HEALTH_LINKS)

        # Print table
        print_summary_table("Device Health Summary", health_results)

        # Save DOCX file
        save_docx_summary(DOCX_DEVICE_HEALTH_PATH, "Device Health Summary", health_results)
        
        # Also save CSV
        save_csv_summary(CSV_DEVICE_HEALTH_PATH, "Device Health Summary", health_results)
        
        print(f"Saved files to:")
        print(f"  DOCX: {DOCX_DEVICE_HEALTH_PATH}")
        print(f"  CSV: {CSV_DEVICE_HEALTH_PATH}")

    finally:
        driver.quit()


if __name__ == "__main__":
    run()
