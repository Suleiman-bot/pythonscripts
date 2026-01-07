"""
connectivity-statistics.py

Reads per-ISP link files (`Glo.txt`, `Dolphin-sec.txt`, `Dolphin-pri.txt`), replaces active/fromOffset/toOffset
with generated constants (same logic as `statistics.py`), visits each metric link with Selenium, polls
for Min/Max values (configurable timeout/poll), prints discoveries and a formatted fixed-width
monitor table per ISP, and writes CSVs.

Requirements from Prompt.txt implemented:
- Table titles and layout per ISP (Glo, Dolphin Sec., Dolphin Pri.) with regions grouped
- Regions: Nigeria (Airtel), Europe (Seabone)  /  US East (TATA), US West (Hurricane)
- Columns: Min., Max., Ref. Link (Ref Link left empty in CSV/table placeholder)
- Rows: Jitter (ms), Latency (ms), Packet Loss (%)
- Polling timeout and printing MIN/MAX as soon as discovered

Note: This script requires Edge + Selenium to run in your environment.
"""

import re
import time
import csv
from datetime import datetime, timedelta
import pytz
import urllib3
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.common.exceptions import TimeoutException

from config import ACCESS_TOKEN,TARGET_DATE

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    Document = None

magic_link = f"https://www.cv-prod-euwest-2.arista.io/api/v1/oauth?invitation={ACCESS_TOKEN}"

# === Active generation function copied from statistics.py ===

def generate_active_for_target_date_local(target_date_str=None):
    """
    Returns 'active' timestamp in milliseconds for the target date's 00:00:00
    in local timezone (Africa/Lagos, UTC+1)
    
    Args:
        target_date_str: Date string in format "M/D/YYYY". If None, uses TARGET_DATE from config
    """
    if target_date_str is None:
        target_date_str = TARGET_DATE
        
    local_tz = pytz.timezone("Africa/Lagos")
    
    # Parse the target date string in format "M/D/YYYY"
    target_date = datetime.strptime(target_date_str, "%m/%d/%Y")
    
    # Localize to timezone without adding extra hours
    target_date_local = local_tz.localize(target_date)
    
    # Convert to UTC timestamp in milliseconds
    active_ts = int(target_date_local.timestamp() * 1000)
    return active_ts

ACTIVE = generate_active_for_target_date_local()
FROM_OFFSET = 1000
TO_OFFSET = 86400000
# === End of active generation function ===


# DOCX output paths
DOCX_GLO = r"C:\Users\SuleimanAbdulsalam\Downloads\GLO.docx"
DOCX_DOL_SEC = r"C:\Users\SuleimanAbdulsalam\Downloads\Dolphin-sec.docx"
DOCX_DOL_PRI = r"C:\Users\SuleimanAbdulsalam\Downloads\Dolphin-pri.docx"
# CSV output paths (kept for backward compatibility and matching table layout)
CSV_GLO = r"C:\Users\SuleimanAbdulsalam\Downloads\GLO.csv"
CSV_DOL_SEC = r"C:\Users\SuleimanAbdulsalam\Downloads\Dolphin-sec.csv"
CSV_DOL_PRI = r"C:\Users\SuleimanAbdulsalam\Downloads\Dolphin-pri.csv"

# Helpers for URL normalization and numeric extraction
_replace_active_re = re.compile(r"active=[^&]*")
_replace_from_re = re.compile(r"fromOffset=[^&]*")
_replace_to_re = re.compile(r"toOffset=[^&]*")
_num_re = re.compile(r"[-+]?[0-9]+(?:[\,\.][0-9]+)*")


def extract_first_number(text):
    """Return first numeric substring or empty string."""
    if not text:
        return ""
    m = _num_re.search(text.replace('\u00a0', ' '))
    if not m:
        return ""
    return m.group(0)


def normalize_url(url: str) -> str:
    """Replace active/fromOffset/toOffset with generated constants."""
    url = _replace_active_re.sub(f"active={ACTIVE}", url)
    url = _replace_from_re.sub(f"fromOffset={FROM_OFFSET}", url)
    url = _replace_to_re.sub(f"toOffset={TO_OFFSET}", url)
    return url


# Polling finder for min and max values; prints discoveries immediately
def find_min_max_for_metric(driver, label="", timeout=12, poll=0.5):
    start = time.time()
    prev_min = None
    prev_max = None

    def _attempt_find():
        try:
            WebDriverWait(driver, 3).until(lambda d: d.execute_script("return document.readyState") == "complete")
        except TimeoutException:
            pass

        min_val = ""
        max_val = ""

        # attempt DOM lookups
        try:
            min_el = driver.find_element(By.XPATH, "//*[contains(translate(text(),'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'min')]")
        except Exception:
            min_el = None
        try:
            max_el = driver.find_element(By.XPATH, "//*[contains(translate(text(),'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'max')]")
        except Exception:
            max_el = None

        if min_el is not None:
            try:
                txt = min_el.text or ""
                if not extract_first_number(txt):
                    try:
                        txt = min_el.find_element(By.XPATH, "..").text
                    except Exception:
                        txt = txt
                min_val = extract_first_number(txt)
            except Exception:
                min_val = ""

        if max_el is not None:
            try:
                txt = max_el.text or ""
                if not extract_first_number(txt):
                    try:
                        txt = max_el.find_element(By.XPATH, "..").text
                    except Exception:
                        txt = txt
                max_val = extract_first_number(txt)
            except Exception:
                max_val = ""

        # fallback to searching body text
        if (not min_val) or (not max_val):
            try:
                page_txt = driver.find_element(By.TAG_NAME, "body").text
            except Exception:
                page_txt = ""
            m_min = re.search(r"min[^0-9\n\r\$\-]*([0-9]+(?:[\,\.][0-9]+)*)", page_txt, re.IGNORECASE)
            m_max = re.search(r"max[^0-9\n\r\$\-]*([0-9]+(?:[\,\.][0-9]+)*)", page_txt, re.IGNORECASE)
            if m_min and not min_val:
                min_val = m_min.group(1)
            if m_max and not max_val:
                max_val = m_max.group(1)

        return min_val, max_val

    while time.time() - start < timeout:
        min_val, max_val = _attempt_find()
        if min_val and min_val != prev_min:
            print(f"Found MIN for {label}: {min_val}") if label else print(f"Found MIN: {min_val}")
            prev_min = min_val
        if max_val and max_val != prev_max:
            print(f"Found MAX for {label}: {max_val}") if label else print(f"Found MAX: {max_val}")
            prev_max = max_val
        if prev_min and prev_max:
            return prev_min, prev_max
        time.sleep(poll)

    return prev_min or "", prev_max or ""


# Collect metrics from a hardcoded dictionary { metric: { region: url } }
def collect_for_isp_dict(driver, isp_name, links_dict, per_page_timeout=12):
    print(f"\n=== Collecting for {isp_name} (hardcoded links) ===")
    results = {}
    for metric, regions in links_dict.items():
        results.setdefault(metric, {})
        for region, url in regions.items():
            normalized = normalize_url(url)
            try:
                driver.get(normalized)
                time.sleep(0.5)
                print(f"Opening {isp_name} - {metric} - {region} ...")
                mn, mx = find_min_max_for_metric(driver, label=f"{isp_name} {metric} {region}", timeout=per_page_timeout)
                results[metric][region] = (mn, mx, normalized)
                print(f"Result {isp_name} - {metric} - {region} -> min: '{mn}', max: '{mx}'")
            except Exception as e:
                print(f"Error {isp_name} {metric} {region}: {e}")
                results[metric][region] = ("", "", normalized)

    return results

# Hardcoded links copied from the .txt files and adapted to use ACTIVE/FROM_OFFSET/TO_OFFSET
GLO_LINKS = {
    "Jitter": {
        "Airtel": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Jitter+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet1%29+to+GLO_INT_to_AIRTEL_NIG%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22JITTER_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet1%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22GLO_INT_to_AIRTEL_NIG%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=jitter",
        "seabone Europe": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Jitter+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet1%29+to+GLO_INT_to_SEABON_ITA%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22JITTER_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet1%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22GLO_INT_to_SEABON_ITA%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=jitter",
        "Tata US East": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Jitter+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet1%29+to+GLO_INT_to_TATA_East_US%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22JITTER_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet1%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22GLO_INT_to_TATA_East_US%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=jitter",
        "Hurricane US west": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Jitter+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet1%29+to+GLO_INT_to_Hurri_West_US%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22JITTER_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet1%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22GLO_INT_to_Hurri_West_US%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=jitter",
    },
    "Latency": {
        "Airtel": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Latency+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet1%29+to+GLO_INT_to_AIRTEL_NIG%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22LATENCY_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet1%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22GLO_INT_to_AIRTEL_NIG%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=latency",
        "seabone Europe": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Latency+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet1%29+to+GLO_INT_to_SEABON_ITA%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22LATENCY_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet1%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22GLO_INT_to_SEABON_ITA%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=latency",
        "Tata US East": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Latency+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet1%29+to+GLO_INT_to_TATA_East_US%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22LATENCY_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet1%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22GLO_INT_to_TATA_East_US%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=latency",
        "Hurricane US west": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Latency+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet1%29+to+GLO_INT_to_Hurri_West_US%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22LATENCY_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet1%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22GLO_INT_to_Hurri_West_US%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=latency",
    },
    "Packet Loss": {
        "Airtel": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Packet+Loss+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet1%29+to+GLO_INT_to_AIRTEL_NIG%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22PACKET_LOSS_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet1%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22GLO_INT_to_AIRTEL_NIG%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=loss",
        "seabone Europe": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Packet+Loss+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet1%29+to+GLO_INT_to_SEABON_ITA%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22PACKET_LOSS_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet1%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22GLO_INT_to_SEABON_ITA%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=loss",
        "Tata US East": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Packet+Loss+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet1%29+to+GLO_INT_to_TATA_East_US%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22PACKET_LOSS_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet1%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22GLO_INT_to_TATA_East_US%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=loss",
        "Hurricane US west": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Packet+Loss+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet1%29+to+GLO_INT_to_Hurri_West_US%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22PACKET_LOSS_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet1%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22GLO_INT_to_Hurri_West_US%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=loss",
    }
}

DOL_SEC_LINKS = {
    "Jitter": {
        "Airtel": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Jitter+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet2%29+to+DOL_SEC_to_AIRTEL_NIG%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22JITTER_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet2%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22DOL_SEC_to_AIRTEL_NIG%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=jitter",
        "seabone Europe": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Jitter+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet2%29+to+DOL_SEC_to_SEABON_ITA%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22JITTER_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet2%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22DOL_SEC_to_SEABON_ITA%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=jitter",
        "Tata US East": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Jitter+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet2%29+to+DOL_SEC_to_TATA_East_US%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22JITTER_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet2%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22DOL_SEC_to_TATA_East_US%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=jitter",
        "Hurricane US west": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Jitter+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet2%29+to+DOL_SEC_to_Hurri_West_US%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22JITTER_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet2%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22DOL_SEC_to_Hurri_West_US%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=jitter",
    },
    "Latency": {
        "Airtel": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Latency+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet2%29+to+DOL_SEC_to_AIRTEL_NIG%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22LATENCY_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet2%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22DOL_SEC_to_AIRTEL_NIG%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=latency",
        "seabone Europe": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Latency+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet2%29+to+DOL_SEC_to_SEABON_ITA%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22LATENCY_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet2%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22DOL_SEC_to_SEABON_ITA%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=latency",
        "Tata US East": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Latency+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet2%29+to+DOL_SEC_to_TATA_East_US%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22LATENCY_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet2%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22DOL_SEC_to_TATA_East_US%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=latency",
        "Hurricane US west": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Latency+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet2%29+to+DOL_SEC_to_Hurri_West_US%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22LATENCY_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet2%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22DOL_SEC_to_Hurri_West_US%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=latency",
    },
    "Packet Loss": {
        "Airtel": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Packet+Loss+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet2%29+to+DOL_SEC_to_AIRTEL_NIG%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22PACKET_LOSS_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet2%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22DOL_SEC_to_AIRTEL_NIG%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=loss",
        "seabone Europe": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Packet+Loss+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet2%29+to+DOL_SEC_to_SEABON_ITA%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22PACKET_LOSS_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet2%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22DOL_SEC_to_SEABON_ITA%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=loss",
        "Tata US East": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Packet+Loss+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet2%29+to+DOL_SEC_to_TATA_East_US%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22PACKET_LOSS_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet2%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22DOL_SEC_to_TATA_East_US%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=loss",
        "Hurricane US west": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Packet+Loss+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet2%29+to+DOL_SEC_to_Hurri_West_US%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22PACKET_LOSS_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet2%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22DOL_SEC_to_Hurri_West_US%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=loss",
    }
}

DOL_PRI_LINKS = {
    "Jitter": {
        "Airtel": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Jitter+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet3%29+to+DOL_PRI_to_AIRTEL_NIG%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22JITTER_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet3%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22DOL_PRI_to_AIRTEL_NIG%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=jitter",
        "seabone Europe": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Jitter+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet3%29+to+DOL_PRI_to_SEABON_ITA%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22JITTER_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet3%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22DOL_PRI_to_SEABON_ITA%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=jitter",
        "Tata US East": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Jitter+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet3%29+to+DOL_PRI_to_TATA_East_US%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22JITTER_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet3%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22DOL_PRI_to_TATA_East_US%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=jitter",
        "Hurricane US west": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Jitter+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet3%29+to+DOL_PRI_to_Hurri_West_US%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22JITTER_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet3%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22DOL_PRI_to_Hurri_West_US%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=jitter",
    },
    "Latency": {
        "Airtel": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Latency+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet3%29+to+DOL_PRI_to_AIRTEL_NIG%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22LATENCY_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet3%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22DOL_PRI_to_AIRTEL_NIG%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=latency",
        "seabone Europe": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Latency+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet3%29+to+DOL_PRI_to_SEABON_ITA%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22LATENCY_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet3%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22DOL_PRI_to_SEABON_ITA%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=latency",
        "Tata US East": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Latency+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet3%29+to+DOL_PRI_to_TATA_East_US%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22LATENCY_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet3%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22DOL_PRI_to_TATA_East_US%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=latency",
        "Hurricane US west": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Latency+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet3%29+to+DOL_PRI_to_Hurri_West_US%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22LATENCY_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet3%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22DOL_PRI_to_Hurri_West_US%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=latency",
    },
    "Packet Loss": {
        "Airtel": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Packet+Loss+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet3%29+to+DOL_PRI_to_AIRTEL_NIG%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22PACKET_LOSS_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet3%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22DOL_PRI_to_AIRTEL_NIG%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=loss",
        "seabone Europe": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Packet+Loss+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet3%29+to+DOL_PRI_to_SEABON_ITA%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22PACKET_LOSS_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet3%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22DOL_PRI_to_SEABON_ITA%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=loss",
        "Tata US East": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Packet+Loss+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet3%29+to+DOL_PRI_to_TATA_East_US%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22PACKET_LOSS_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet3%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22DOL_PRI_to_TATA_East_US%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=loss",
        "Hurricane US west": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/connectivity?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22alternativeTitle%22%3A%22Packet+Loss+for+KASI-LOS5-R201-BG01+%28inet+%2F+Ethernet3%29+to+DOL_PRI_to_Hurri_West_US%22%2C%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22PACKET_LOSS_VRF%22%2C%22metricParams%22%3A%7B%22deviceId%22%3A%22JPE20050335%22%2C%22hostVrfIntf%22%3A%22Ethernet3%22%2C%22hostVrfPair%22%3A%7B%22hostName%22%3A%22DOL_PRI_to_Hurri_West_US%22%2C%22vrfName%22%3A%22inet%22%7D%7D%7D&modalPanel=STATISTICS&metric=loss",
    }
}

# Print fixed-width tables per prompt
def print_isp_table(title, results):
    # Column widths
    metric_w = 18
    val_w = 10
    ref_w = 28

    # Title on its own line
    print(title)

    # Section 1: Nigeria (Airtel) | Europe (Seabone)
    left_label = "Nigeria (Airtel)"
    right_label = "Europe (Seabone)"
    header = "".ljust(metric_w) + f"  {left_label.center(val_w*2 + ref_w + 4)}  {right_label.center(val_w*2 + ref_w + 4)}"
    print(header)
    sub = "".ljust(metric_w) + "  " + "Min.".center(val_w) + " " + "Max.".center(val_w) + " " + "Ref. Link".center(ref_w) + "  " + "Min.".center(val_w) + " " + "Max.".center(val_w) + " " + "Ref. Link".center(ref_w)
    print(sub)
    print("-" * (metric_w + (val_w*4) + (ref_w*2) + 12))

    metrics = ["Jitter (ms)", "Latency (ms)", "Packet Loss (%)"]
    for metric in metrics:
        key = metric.split(" (")[0]
        left = results.get(key, {}).get("Airtel", ("", "", ""))
        right = results.get(key, {}).get("seabone Europe", ("", "", ""))
        print(f"{metric.ljust(metric_w)}  {str(left[0]).center(val_w)} {str(left[1]).center(val_w)} {''.center(ref_w)}  {str(right[0]).center(val_w)} {str(right[1]).center(val_w)} {''.center(ref_w)}")

    print()

    # Section 2: US East (TATA) | US West (Hurricane)
    left_label = "US East (TATA)"
    right_label = "US West (Hurricane)"
    header = "".ljust(metric_w) + f"  {left_label.center(val_w*2 + ref_w + 4)}  {right_label.center(val_w*2 + ref_w + 4)}"
    print(header)
    print(sub)
    print("-" * (metric_w + (val_w*4) + (ref_w*2) + 12))

    for metric in metrics:
        key = metric.split(" (")[0]
        left = results.get(key, {}).get("Tata US East", ("", "", ""))
        right = results.get(key, {}).get("Hurricane US west", ("", "", ""))
        print(f"{metric.ljust(metric_w)}  {str(left[0]).center(val_w)} {str(left[1]).center(val_w)} {''.center(ref_w)}  {str(right[0]).center(val_w)} {str(right[1]).center(val_w)} {''.center(ref_w)}")

    print()


def save_isp_docx(path, title, results):
    if Document is None:
        print("python-docx not installed. Install with: pip install python-docx")
        return
    doc = Document()
    doc.add_heading(title, level=2)

    # Create table with initial rows for top section header and subheader
    table = doc.add_table(rows=2, cols=7)
    table.style = 'Table Grid'

    # Top section header
    hdr_top = table.rows[0].cells
    hdr_top[0].text = ''
    hdr_top[1].merge(hdr_top[3]).text = 'Nigeria (Airtel)'
    hdr_top[4].merge(hdr_top[6]).text = 'Europe (Seabone)'

    # Top subheader
    sub_top = table.rows[1].cells
    sub_top[0].text = ''
    sub_top[1].text = 'Min'
    sub_top[2].text = 'Max'
    sub_top[3].text = 'Ref. Link'
    sub_top[4].text = 'Min'
    sub_top[5].text = 'Max'
    sub_top[6].text = 'Ref. Link'

    # Add top section data rows
    metrics = ['Jitter (ms)', 'Latency (ms)', 'Packet Loss (%)']
    for metric in metrics:
        key = metric.split(' (')[0]
        row = table.add_row().cells
        row[0].text = metric
        left = results.get(key, {}).get('Airtel', ('', '', ''))
        right = results.get(key, {}).get('seabone Europe', ('', '', ''))
        row[1].text = left[0]
        row[2].text = left[1]
        row[3].text = ''
        row[4].text = right[0]
        row[5].text = right[1]
        row[6].text = ''

    # Add bottom section header
    hdr_bottom_row = table.add_row()
    hdr_bottom = hdr_bottom_row.cells
    hdr_bottom[0].text = ''
    hdr_bottom[1].merge(hdr_bottom[3]).text = 'US East (TATA)'
    hdr_bottom[4].merge(hdr_bottom[6]).text = 'US West (Hurricane)'

    # Add bottom subheader
    sub_bottom_row = table.add_row()
    sub_bottom = sub_bottom_row.cells
    sub_bottom[0].text = ''
    sub_bottom[1].text = 'Min'
    sub_bottom[2].text = 'Max'
    sub_bottom[3].text = 'Ref. Link'
    sub_bottom[4].text = 'Min'
    sub_bottom[5].text = 'Max'
    sub_bottom[6].text = 'Ref. Link'

    # Add bottom section data rows
    for metric in metrics:
        key = metric.split(' (')[0]
        row = table.add_row().cells
        row[0].text = metric
        left = results.get(key, {}).get('Tata US East', ('', '', ''))
        right = results.get(key, {}).get('Hurricane US west', ('', '', ''))
        row[1].text = left[0]
        row[2].text = left[1]
        row[3].text = ''
        row[4].text = right[0]
        row[5].text = right[1]
        row[6].text = ''

    # Bold header rows (top: 0-1, bottom: current index after top data)
    # After top header/sub + 3 data = rows 0-4, then hdr_bottom 5, sub_bottom 6
    header_rows = [table.rows[0], table.rows[1], table.rows[5], table.rows[6]]
    for row in header_rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True

    doc.save(path)
    print(f"Saved DOCX to: {path}")


# Save CSV that mirrors the two DOCX tables in a single CSV layout
def save_isp_csv(path, title, results):
    with open(path, 'w', newline='', encoding='utf-8') as fh:
        w = csv.writer(fh)
        w.writerow([title])
        # Combined header for both regional groupings
        header = ["Metric",
                  "Airtel Min", "Airtel Max", "Airtel Ref",
                  "Seabone Min", "Seabone Max", "Seabone Ref",
                  "Tata Min", "Tata Max", "Tata Ref",
                  "Hurricane Min", "Hurricane Max", "Hurricane Ref"]
        w.writerow(header)
        for metric in ['Jitter', 'Latency', 'Packet Loss']:
            left = results.get(metric, {}).get('Airtel', ('', '', ''))
            seabone = results.get(metric, {}).get('seabone Europe', ('', '', ''))
            tata = results.get(metric, {}).get('Tata US East', ('', '', ''))
            hur = results.get(metric, {}).get('Hurricane US west', ('', '', ''))
            row = [metric,
                   left[0], left[1], '',
                   seabone[0], seabone[1], '',
                   tata[0], tata[1], '',
                   hur[0], hur[1], '']
            w.writerow(row)
    print(f"Saved CSV to: {path}")


def run():
    print("🌐 Opening Selenium with Edge profile...")
    options = webdriver.EdgeOptions()
    options.add_argument("--start-maximized")
    options.add_argument(r"--user-data-dir=C:\Users\SuleimanAbdulsalam\AppData\Local\Microsoft\Edge\User Data\SeleniumProfile")
    options.add_argument("profile-directory=Default")
    driver = webdriver.Edge(options=options)

    try:
        print("Opening CVaaS login link...")
        driver.get(magic_link)
        try:
            WebDriverWait(driver, 10).until(lambda d: d.current_url != magic_link or d.execute_script("return document.readyState") == "complete")
        except TimeoutException:
            pass

        glo_results = collect_for_isp_dict(driver, "Glo", GLO_LINKS, per_page_timeout=12)
        dol_sec_results = collect_for_isp_dict(driver, "Dolphin Sec.", DOL_SEC_LINKS, per_page_timeout=12)
        dol_pri_results = collect_for_isp_dict(driver, "Dolphin Pri.", DOL_PRI_LINKS, per_page_timeout=12)

        # Print tables per prompt
        print("\nGlo Traffic Monitor")
        print_isp_table("Glo Traffic Monitor", glo_results)
        print("Dolphin Sec. Traffic Monitor")
        print_isp_table("Dolphin Sec. Traffic Monitor", dol_sec_results)
        print("Dolphin Pri. Traffic Monitor")
        print_isp_table("Dolphin Pri. Traffic Monitor", dol_pri_results)

        # Save DOCX files
        save_isp_docx(DOCX_GLO, "Glo Traffic Monitor", glo_results)
        save_isp_docx(DOCX_DOL_SEC, "Dolphin Sec. Traffic Monitor", dol_sec_results)
        save_isp_docx(DOCX_DOL_PRI, "Dolphin Pri. Traffic Monitor", dol_pri_results)
        # Also save matching CSV tables
        save_isp_csv(CSV_GLO, "Glo Traffic Monitor", glo_results)
        save_isp_csv(CSV_DOL_SEC, "Dolphin Sec. Traffic Monitor", dol_sec_results)
        save_isp_csv(CSV_DOL_PRI, "Dolphin Pri. Traffic Monitor", dol_pri_results)
        print(f"Saved DOCXs to: {DOCX_GLO}, {DOCX_DOL_SEC}, {DOCX_DOL_PRI}")
        print(f"Saved CSVs to: {CSV_GLO}, {CSV_DOL_SEC}, {CSV_DOL_PRI}")

    finally:
        driver.quit()


if __name__ == "__main__":
    run()