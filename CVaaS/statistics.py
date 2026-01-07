from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, StaleElementReferenceException
import time
import re
import csv
from datetime import datetime, timedelta
import pytz
import urllib3
from config import ACCESS_TOKEN, TARGET_DATE

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    Document = None

# === Construct magic login link (same approach as export.py) ===
magic_link = f"https://www.cv-prod-euwest-2.arista.io/api/v1/oauth?invitation={ACCESS_TOKEN}"


def generate_active_for_target_date_local(target_date_str=None):
    """
    Returns 'active' timestamp in milliseconds for the target date's 00:00:00
    in local timezone (Africa/Lagos, UTC+1)
    
    Args:
        target_date_str: Date string in format "M/D/YYYY". If None, uses TARGET_DATE from config
    """
    if target_date_str is None:
        target_date_str = TARGET_DATE
        
    local_tz = pytz.timezone("Africa/Lagos")
    
    # Parse the target date string in format "M/D/YYYY"
    target_date = datetime.strptime(target_date_str, "%m/%d/%Y")
    
    # Localize to timezone without adding extra hours
    target_date_local = local_tz.localize(target_date)
    
    # Convert to UTC timestamp in milliseconds
    active_ts = int(target_date_local.timestamp() * 1000)
    return active_ts

ACTIVE = generate_active_for_target_date_local()
FROM_OFFSET = 1000
TO_OFFSET = 86400000
# === End of active generation function ===


# --- Links (constants) but include active/from/to via variables above ---
BITRATE_LINKS = {
    "Glo": {
        "inbound": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/ethernet-stats/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22INTERFACE_IN_BITRATE%22%2C%22metricParams%22%3A%7B%22aggregationIntervalOverride%22%3A%221m%22%2C%22deviceId%22%3A%22JPE20050335%22%2C%22intf%22%3A%22Ethernet1%22%7D%7D&modalPanel=STATISTICS",
        "outbound": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/ethernet-stats/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22INTERFACE_OUT_BITRATE%22%2C%22metricParams%22%3A%7B%22aggregationIntervalOverride%22%3A%221m%22%2C%22deviceId%22%3A%22JPE20050335%22%2C%22intf%22%3A%22Ethernet1%22%7D%7D&modalPanel=STATISTICS",
    },
    "Dolphin Sec.": {
        "inbound": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/ethernet-stats/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22INTERFACE_IN_BITRATE%22%2C%22metricParams%22%3A%7B%22aggregationIntervalOverride%22%3A%221m%22%2C%22deviceId%22%3A%22JPE20050335%22%2C%22intf%22%3A%22Ethernet2%22%7D%7D&modalPanel=STATISTICS",
        "outbound": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/ethernet-stats/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22INTERFACE_OUT_BITRATE%22%2C%22metricParams%22%3A%7B%22aggregationIntervalOverride%22%3A%221m%22%2C%22deviceId%22%3A%22JPE20050335%22%2C%22intf%22%3A%22Ethernet2%22%7D%7D&modalPanel=STATISTICS",
    },
    "Dolphin Pri.": {
        "inbound": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/ethernet-stats/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22INTERFACE_IN_BITRATE%22%2C%22metricParams%22%3A%7B%22aggregationIntervalOverride%22%3A%221m%22%2C%22deviceId%22%3A%22JPE20050335%22%2C%22intf%22%3A%22Ethernet3%22%7D%7D&modalPanel=STATISTICS",
        "outbound": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/ethernet-stats/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22INTERFACE_OUT_BITRATE%22%2C%22metricParams%22%3A%7B%22aggregationIntervalOverride%22%3A%221m%22%2C%22deviceId%22%3A%22JPE20050335%22%2C%22intf%22%3A%22Ethernet3%22%7D%7D&modalPanel=STATISTICS",
    },
}

PACKET_LINKS = {
    "Glo": {
        "inbound": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/ethernet-stats/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22INTERFACE_IN_UCAST_RATE%22%2C%22metricParams%22%3A%7B%22aggregationIntervalOverride%22%3A%221m%22%2C%22deviceId%22%3A%22JPE20050335%22%2C%22intf%22%3A%22Ethernet1%22%7D%7D&modalPanel=STATISTICS",
        "outbound": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/ethernet-stats/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22INTERFACE_OUT_UCAST_RATE%22%2C%22metricParams%22%3A%7B%22aggregationIntervalOverride%22%3A%221m%22%2C%22deviceId%22%3A%22JPE20050335%22%2C%22intf%22%3A%22Ethernet1%22%7D%7D&modalPanel=STATISTICS",
    },
    "Dolphin Sec.": {
        "inbound": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/ethernet-stats/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22INTERFACE_IN_UCAST_RATE%22%2C%22metricParams%22%3A%7B%22aggregationIntervalOverride%22%3A%221m%22%2C%22deviceId%22%3A%22JPE20050335%22%2C%22intf%22%3A%22Ethernet2%22%7D%7D&modalPanel=STATISTICS",
        "outbound": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/ethernet-stats/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22INTERFACE_OUT_UCAST_RATE%22%2C%22metricParams%22%3A%7B%22aggregationIntervalOverride%22%3A%221m%22%2C%22deviceId%22%3A%22JPE20050335%22%2C%22intf%22%3A%22Ethernet2%22%7D%7D&modalPanel=STATISTICS",
    },
    "Dolphin Pri.": {
        "inbound": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/ethernet-stats/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22INTERFACE_IN_UCAST_RATE%22%2C%22metricParams%22%3A%7B%22aggregationIntervalOverride%22%3A%221m%22%2C%22deviceId%22%3A%22JPE20050335%22%2C%22intf%22%3A%22Ethernet3%22%7D%7D&modalPanel=STATISTICS",
        "outbound": f"https://www.cv-prod-euwest-2.arista.io/cv/devices/ethernet-stats/JPE20050335?active={ACTIVE}&fromOffset={FROM_OFFSET}&toOffset={TO_OFFSET}&modal=true&modalParams=%7B%22datasetId%22%3A%22JPE20050335%22%2C%22metricKey%22%3A%22INTERFACE_OUT_UCAST_RATE%22%2C%22metricParams%22%3A%7B%22aggregationIntervalOverride%22%3A%221m%22%2C%22deviceId%22%3A%22JPE20050335%22%2C%22intf%22%3A%22Ethernet3%22%7D%7D&modalPanel=STATISTICS",
    },
}

DOCX_BITRATE_PATH = r"C:\Users\SuleimanAbdulsalam\Downloads\bitrate-summary.docx"
DOCX_PACKET_PATH = r"C:\Users\SuleimanAbdulsalam\Downloads\packetrate-summary.docx"
CSV_BITRATE_PATH = r"C:\Users\SuleimanAbdulsalam\Downloads\bitrate-summary.csv"
CSV_PACKET_PATH = r"C:\Users\SuleimanAbdulsalam\Downloads\packetrate-summary.csv"


# Utility: try to parse a number (integers/floats) from a text blob
_num_re = re.compile(r"[-+]?[0-9]+(?:[\,\.][0-9]+)*")

def extract_first_number(text):
    if not text:
        return ""
    m = _num_re.search(text.replace('\u00a0', ' '))
    if not m:
        return ""
    val = m.group(0)
    return val


# Try to find min and max on a page using heuristics
# This version will poll for up to `timeout` seconds and prints values as they are discovered.
def find_min_max_for_metric(driver, label="", timeout=12, poll=0.5):
    start = time.time()
    prev_min = None
    prev_max = None

    def _attempt_find():
        # Ensure page loaded
        try:
            WebDriverWait(driver, 3).until(lambda d: d.execute_script("return document.readyState") == "complete")
        except TimeoutException:
            pass

        # First: look for elements with "min" and "max" text
        min_el = None
        max_el = None
        try:
            min_el = driver.find_element(By.XPATH, "//*[contains(translate(text(),'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'min')]")
        except Exception:
            min_el = None
        try:
            max_el = driver.find_element(By.XPATH, "//*[contains(translate(text(),'ABCDEFGHIJKLMNOPQRSTUVWXYZ','abcdefghijklmnopqrstuvwxyz'),'max')]")
        except Exception:
            max_el = None

        min_val = ""
        max_val = ""

        if min_el is not None:
            try:
                txt = min_el.text or ""
                if not extract_first_number(txt):
                    try:
                        parent_txt = min_el.find_element(By.XPATH, "..").text
                    except Exception:
                        parent_txt = ""
                    txt = parent_txt
                min_val = extract_first_number(txt)
            except Exception:
                min_val = ""

        if max_el is not None:
            try:
                txt = max_el.text or ""
                if not extract_first_number(txt):
                    try:
                        parent_txt = max_el.find_element(By.XPATH, "..").text
                    except Exception:
                        parent_txt = ""
                    txt = parent_txt
                max_val = extract_first_number(txt)
            except Exception:
                max_val = ""

        # If either missing, try to look for pattern like "Min: 123" or "min 123" across page text
        if not min_val or not max_val:
            try:
                page_txt = driver.find_element(By.TAG_NAME, "body").text
            except Exception:
                page_txt = ""
            m_min = re.search(r"min[^0-9\n\r\$\-]*([0-9]+(?:[\,\.][0-9]+)*)", page_txt, re.IGNORECASE)
            m_max = re.search(r"max[^0-9\n\r\$\-]*([0-9]+(?:[\,\.][0-9]+)*)", page_txt, re.IGNORECASE)
            if m_min and not min_val:
                min_val = m_min.group(1)
            if m_max and not max_val:
                max_val = m_max.group(1)

        return min_val, max_val

    # Poll until both values found or timeout
    while time.time() - start < timeout:
        min_val, max_val = _attempt_find()

        # Print any newly-discovered values immediately
        if min_val and min_val != prev_min:
            if label:
                print(f"Found MIN for {label}: {min_val}")
            else:
                print(f"Found MIN: {min_val}")
            prev_min = min_val
        if max_val and max_val != prev_max:
            if label:
                print(f"Found MAX for {label}: {max_val}")
            else:
                print(f"Found MAX: {max_val}")
            prev_max = max_val

        # If both found, return them
        if min_val and max_val:
            return min_val, max_val

        time.sleep(poll)

    # Final attempt and return whatever was found
    return prev_min or "", prev_max or ""


# Open each URL and extract metric values
def collect_metrics(driver, groups, per_page_timeout=12):
    results = {}
    wait = WebDriverWait(driver, 10)
    for name, urls in groups.items():
        results[name] = {"inbound_min": "", "inbound_max": "", "outbound_min": "", "outbound_max": ""}

        # inbound
        try:
            driver.get(urls["inbound"])
            time.sleep(0.5)
            print(f"Opening {name} inbound URL...")
            imin, imax = find_min_max_for_metric(driver, label=f"{name} inbound", timeout=per_page_timeout)
            results[name]["inbound_min"] = imin
            results[name]["inbound_max"] = imax
            print(f"Result {name} inbound -> min: '{imin}', max: '{imax}'")
        except Exception as e:
            print(f"Error collecting inbound for {name}: {e}")
            results[name]["inbound_min"] = ""
            results[name]["inbound_max"] = ""

        # outbound
        try:
            driver.get(urls["outbound"])
            time.sleep(0.5)
            print(f"Opening {name} outbound URL...")
            omin, omax = find_min_max_for_metric(driver, label=f"{name} outbound", timeout=per_page_timeout)
            results[name]["outbound_min"] = omin
            results[name]["outbound_max"] = omax
            print(f"Result {name} outbound -> min: '{omin}', max: '{omax}'")
        except Exception as e:
            print(f"Error collecting outbound for {name}: {e}")
            results[name]["outbound_min"] = ""
            results[name]["outbound_max"] = ""

    return results


# Print fixed-width table as requested
def print_summary_table(title, results):
    # Column widths
    name_w = 15
    val_w = 12

    total_w = name_w + val_w * 4 + 6  # paddings

    # Title
    print(title)

    # Header with grouped columns
    name_col = "Link name".ljust(name_w)
    inbound_header = "Inbound".center(val_w * 2 + 1)
    outbound_header = "Outbound".center(val_w * 2 + 1)
    print(f"{name_col}  {inbound_header}  {outbound_header}")

    # Sub-headers
    sub = "".ljust(name_w) + "  " + "Min".center(val_w) + " " + "Max".center(val_w) + "  " + "Min".center(val_w) + " " + "Max".center(val_w)
    print(sub)

    # Separator
    print("-" * (name_w + val_w * 4 + 6))

    # Rows
    for name in ["Glo", "Dolphin Sec.", "Dolphin Pri."]:
        row = results.get(name, {})
        a_min = row.get("inbound_min", "") or ""
        a_max = row.get("inbound_max", "") or ""
        b_min = row.get("outbound_min", "") or ""
        b_max = row.get("outbound_max", "") or ""
        print(f"{name.ljust(name_w)}  {a_min.center(val_w)} {a_max.center(val_w)}  {b_min.center(val_w)} {b_max.center(val_w)}")


# Save DOCX
def save_docx_summary(path, title, results, unit=None):
    if Document is None:
        print("python-docx not installed. Install with: pip install python-docx")
        return
    doc = Document()
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Table Grid'
    # Header row: merge inbound/outbound
    hdr = table.rows[0].cells
    hdr[0].text = 'Link name'
    hdr[1].merge(hdr[2]).text = 'Inbound'
    hdr[3].merge(hdr[4]).text = 'Outbound'
    # Subheader
    sub = table.rows[1].cells
    sub[0].text = ''
    sub[1].text = 'Min'
    sub[2].text = 'Max'
    sub[3].text = 'Min'
    sub[4].text = 'Max'
    # Bold headers
    for cell in hdr + sub:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    # Rows
    for name in ['Glo', 'Dolphin Sec.', 'Dolphin Pri.']:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        a_min = results.get(name, {}).get('inbound_min','') or ''
        a_max = results.get(name, {}).get('inbound_max','') or ''
        b_min = results.get(name, {}).get('outbound_min','') or ''
        b_max = results.get(name, {}).get('outbound_max','') or ''
        row_cells[1].text = f"{a_min} {unit}" if unit and a_min else a_min
        row_cells[2].text = f"{a_max} {unit}" if unit and a_max else a_max
        row_cells[3].text = f"{b_min} {unit}" if unit and b_min else b_min
        row_cells[4].text = f"{b_max} {unit}" if unit and b_max else b_max
    doc.save(path)
    print(f"Saved DOCX to: {path}")


# Save CSV with same layout as DOCX table
def save_csv_summary(path, title, results, unit=None):
    with open(path, "w", newline='', encoding='utf-8') as fh:
        w = csv.writer(fh)
        # Header rows that mirror the DOCX layout
        w.writerow([title])
        w.writerow(["Link name", "Inbound Min", "Inbound Max", "Outbound Min", "Outbound Max"])
        for name in ['Glo', 'Dolphin Sec.', 'Dolphin Pri.']:
            row = results.get(name, {})
            a_min = row.get('inbound_min','') or ''
            a_max = row.get('inbound_max','') or ''
            b_min = row.get('outbound_min','') or ''
            b_max = row.get('outbound_max','') or ''
            if unit:
                if a_min: a_min = f"{a_min} {unit}"
                if a_max: a_max = f"{a_max} {unit}"
                if b_min: b_min = f"{b_min} {unit}"
                if b_max: b_max = f"{b_max} {unit}"
            w.writerow([name, a_min, a_max, b_min, b_max])
    print(f"Saved CSV to: {path}")


def run():
    print("🌐 Opening Selenium with Edge profile...")
    options = webdriver.EdgeOptions()
    options.add_argument("--start-maximized")
    options.add_argument(r"--user-data-dir=C:\Users\SuleimanAbdulsalam\AppData\Local\Microsoft\Edge\User Data\SeleniumProfile")
    options.add_argument("profile-directory=Default")
    driver = webdriver.Edge(options=options)

    try:
        print("Opening CVaaS login link...")
        driver.get(magic_link)
        try:
            WebDriverWait(driver, 10).until(lambda d: d.current_url != magic_link or d.execute_script("return document.readyState") == "complete")
        except TimeoutException:
            pass

        print("Collecting bitrate metrics...")
        bitrate_results = collect_metrics(driver, BITRATE_LINKS)

        print("Collecting packet rate metrics...")
        packet_results = collect_metrics(driver, PACKET_LINKS)

        # Print tables
        print_summary_table("Bitrate Summary", bitrate_results)
        print()
        print_summary_table("Packet rate Summary", packet_results)

        # Save DOCX files
        save_docx_summary(DOCX_BITRATE_PATH, "Bitrate Summary", bitrate_results, unit="Mbps")
        save_docx_summary(DOCX_PACKET_PATH, "Packet rate Summary", packet_results, unit="kpps")
        # Also save matching CSV tables
        save_csv_summary(CSV_BITRATE_PATH, "Bitrate Summary", bitrate_results, unit="Mbps")
        save_csv_summary(CSV_PACKET_PATH, "Packet rate Summary", packet_results, unit="kpps")
        print(f"Saved DOCXs to: {DOCX_BITRATE_PATH} and {DOCX_PACKET_PATH}")
        print(f"Saved CSVs to: {CSV_BITRATE_PATH} and {CSV_PACKET_PATH}")

    finally:
        driver.quit()


if __name__ == "__main__":
    run()